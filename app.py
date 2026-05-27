import os, tempfile, subprocess, platform
from pathlib import Path
from flask import Flask, render_template, request, send_file, jsonify
import openpyxl
from docx import Document

SYSTEM = platform.system()

if SYSTEM == 'Windows':
    import win32com.client
    import pythoncom

app = Flask(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_FILE = os.path.join(BASE_DIR, 'Data TKA SMP Mumtaza.xlsx')
TEMPLATE_FILE = os.path.join(BASE_DIR, 'template.docx')

MONTHS = ['Januari','Februari','Maret','April','Mei','Juni',
          'Juli','Agustus','September','Oktober','November','Desember']

def load_data():
    wb = openpyxl.load_workbook(DATA_FILE)
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    data = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        record = dict(zip(headers, row))
        val = record.get('ddmmyyyy')
        if isinstance(val, (float, int)):
            record['ddmmyyyy'] = str(int(val))
        else:
            record['ddmmyyyy'] = str(val).strip() if val else ''
        record['NISN'] = str(record.get('NISN', '')).strip()
        data.append(record)
    return data

DATA = load_data()

def find_student(nisn, date):
    for r in DATA:
        if r['NISN'] == nisn and r['ddmmyyyy'] == date:
            return r
    return None

def format_date(s):
    s = str(s)
    if len(s) == 8 and s.isdigit():
        d, m, y = s[:2], s[2:4], s[4:]
        mn = MONTHS[int(m)-1] if 1 <= int(m) <= 12 else m
        return f"{d} {mn} {y}"
    return s

def clean_num(v):
    if isinstance(v, (int, float)):
        return str(int(v))
    return str(v).strip() if v else ''

def replace_in_doc(doc, context):
    def _replace(p):
        full = p.text
        new = full
        ok = False
        for k, v in context.items():
            for pat in [f'<<{k} >>', f'<<{k}>>']:
                if pat in new:
                    new = new.replace(pat, v)
                    ok = True
        if ok and p.runs:
            for run in p.runs:
                run.text = ''
            p.runs[0].text = new
    for p in doc.paragraphs:
        _replace(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace(p)

def fill_template(record):
    doc = Document(TEMPLATE_FILE)
    context = {
        'Nama': str(record.get('Nama', '')),
        'NISN': str(record.get('NISN', '')),
        'Nomor_Peserta': str(record.get('Nomor_Peserta', '')),
        'Tempat_Tanggal_Lahir': format_date(record.get('ddmmyyyy', '')),
        'Matematika_Score': clean_num(record.get('Matematika_Score')),
        'Matematika_Rating': str(record.get('Matematika_Rating', '')),
        'Bahasa_Indonesia_Score': clean_num(record.get('Bahasa_Indonesia_Score')),
        'Bahasa_Indonesia_Rating': str(record.get('Bahasa_Indonesia_Rating', '')),
    }
    replace_in_doc(doc, context)
    tmp = tempfile.mkdtemp()
    out = os.path.join(tmp, 'hasil_tka.docx')
    doc.save(out)
    return out, tmp

def docx_to_pdf(docx_path, pdf_path):
    if SYSTEM == 'Windows':
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        try:
            doc = word.Documents.Open(str(docx_path))
            doc.SaveAs(str(pdf_path), FileFormat=17)
            doc.Close()
        finally:
            word.Quit()
            pythoncom.CoUninitialize()
    else:
        outdir = str(Path(pdf_path).parent)
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf',
             '--outdir', outdir, str(docx_path)],
            check=True, timeout=60, capture_output=True
        )
    return pdf_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/lookup', methods=['POST'])
def lookup():
    nisn = request.form.get('nisn', '').strip()
    date = request.form.get('date', '').strip()
    student = find_student(nisn, date)
    if student:
        return jsonify({
            'found': True,
            'data': {
                'nama': str(student.get('Nama', '')),
                'nomor_peserta': str(student.get('Nomor_Peserta', '')),
                'nisn': str(student.get('NISN', '')),
                'matematika_score': clean_num(student.get('Matematika_Score')),
                'matematika_rating': str(student.get('Matematika_Rating', '')),
                'bindo_score': clean_num(student.get('Bahasa_Indonesia_Score')),
                'bindo_rating': str(student.get('Bahasa_Indonesia_Rating', '')),
            }
        })
    return jsonify({'found': False})

@app.route('/generate', methods=['POST'])
def generate():
    nisn = request.form.get('nisn', '').strip()
    date = request.form.get('date', '').strip()
    student = find_student(nisn, date)
    if not student:
        return 'Data tidak ditemukan', 404

    docx_path, tmp_dir = fill_template(student)
    pdf_path = os.path.join(tmp_dir, 'hasil_tka.pdf')

    try:
        docx_to_pdf(docx_path, pdf_path)
        return send_file(pdf_path, mimetype='application/pdf')
    except Exception as e:
        print(f"PDF conversion failed: {e}")
        return send_file(
            docx_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='hasil_tka.docx'
        )

if __name__ == '__main__':
    app.run(debug=True, port=5000)
